from docx import Document
from docx.shared import Inches

document = Document('test.docx')

print (document.paragraphs)
width=200
height=200
target="<placeholder"
image_file="F:\\PycharmProjects\\finalproject\\PDF Reporting\\image.jpg"

# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

for table in document.tables:
	for row in table.rows:
		for cell in row.cells:
			print(cell.text)
			if target in cell.text:
				mycell=cell.text
				print(str(mycell))
				arr=(str(mycell)).split('|')
				# print(arr)
			#	print(int(arr[2].split('>')[0]))
				cell.text=""
				run = cell.paragraphs[0].add_run().add_picture(image_file,width=Inches(1),height=Inches(1))
				print(arr[2])
				print(arr[1])
				inline_shape = run
				#document.save('test1.docx')
				break
for para in document.paragraphs:
	if target in para.text:
		print(para.text)
		mycell=para.text.split('|')
		para.text=""
		para.add_run().add_picture(image_file,width=Inches(1),height=Inches(1))
		document.save('test1.docx')
