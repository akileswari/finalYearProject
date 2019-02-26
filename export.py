import os
import sys
from docx import Document
from docx.shared import Inches
import win32com.client as client

def template():
    document = Document('F:\\PycharmProjects\\finalproject\\PDF Reporting\\test.docx')
    print(document.paragraphs)
    width = 200
    height = 200
    target = "<placeholder"
    image_file = "F:\\Downloads\\a.jpg"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                if target in cell.text:
                    mycell = cell.text
                    print(str(mycell))
                    arr = (str(mycell)).split('|')
                    # print(arr)
                    #	print(int(arr[2].split('>')[0]))
                    cell.text = ""
                    run = cell.paragraphs[0].add_run().add_picture(image_file, width=Inches(1), height=Inches(1))
                    print(arr[2])
                    print(arr[1])
                    inline_shape = run
                    # document.save('test1.docx')
                    break
    for para in document.paragraphs:
        if target in para.text:
            print(para.text)
            mycell = para.text.split('|')
            para.text = ""
            para.add_run().add_picture(image_file, width=Inches(1), height=Inches(1))
            document.save('test1.docx')

def covx_to_pdf():
    """Convert a Word .docx to PDF"""
    infile = "F:\\PycharmProjects\\finalproject\\test1.docx"
    outfile = "F:\\PycharmProjects\\finalproject\\out1.pdf"
    wdFormatPDF = 17
    word = client.DispatchEx("Word.Application")
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

