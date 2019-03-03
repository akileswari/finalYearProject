import requests
import pandas as pd
from flask import Flask, render_template, request, session, json, send_file, flash, redirect, url_for,Response
import FetchData as rate
import NewsFeed as feed
import DollarToINR as cd
import PDFCreation as docConv
import os
fname=""
from docx import Document
from docx.shared import Inches
app = Flask(__name__)
app.secret_key="123"
app.config['CACHE_TYPE']="null"

@app.route('/')
def home():
    print("home")
    return render_template('home.html')
@app.route('/dashboard')
def dashboard():
    al,a_d,cu,c_d,ir,i_d,ni,ni_d,d_c,d_a,d_i,d_n,val_c,d_c_lis,val_a,d_a_lis,val_i,d_i_lis,val_n,d_n_lis=rate.curr_rate()
    al_f=feed.news_feed("aluminium")
    cu_f = feed.news_feed("copper")
    ir_f = feed.news_feed("iron")
    ni_f = feed.news_feed("gold")
    inr=cd.indian_rs()
    return render_template('dashb.html',al=al,cu=cu,ir=ir,ni=ni,a_d=a_d,c_d=c_d,i_d=i_d,ni_d=ni_d,al_f=al_f,cu_f=cu_f,ir_f=ir_f,ni_f=ni_f,d_c=d_c,
                           d_a=d_a,d_i=d_i,d_n=d_n,val_c=val_c,d_c_lis=d_c_lis,val_a=val_a,d_a_lis=d_a_lis,val_i=val_i,d_i_lis=d_i_lis,val_n=val_n,
                           d_n_lis=d_n_lis,inr=inr)

@app.route('/download',methods=['POST'])
def downloadFile ():
    print("hello2")
    print(request.method)
    print(request.files['file'])
    f=request.files['file']
    #f=request.args.get("file")
    #print (f)
    #For windows you need to use drive name [ex: F:/Example.pdf]
    # document=Document()
    # document.add_paragraph('A plain paragraph having some ')
    # document.add_picture(f,width=Inches(2.0),height=Inches(2.0))
    # document.save("demo1.docx")
    docConv.template(f)
    fname=f.filename
    pdfDownload = open("F:\\PycharmProjects\\finalproject\\static\\"+fname+".pdf", 'rb').read()
    return Response(
        pdfDownload,mimetype="application/pdf",headers={"Content-disposition":"attachment; filename=output.pdf"})
    #return redirect(url_for("downloadpdf?d=succ"))
    #return send_file("F:\\PycharmProjects\\finalproject\\test1.docx", mimetype='application/*', as_attachment=True,attachment_filename="output.docx")

@app.route('/downloadpdf')
def downloadpdf():
    print(request.method)
    fname=request.args['fn']
    # response=send_file("F:\\PycharmProjects\\finalproject\\output.pdf", mimetype='application/*', as_attachment=True,attachment_filename="output.pdf")
    # response.cache_control.max_age=60*2

    #return response
    return send_file("F:\\PycharmProjects\\finalproject\\static\\"+fname+".pdf", mimetype='application/*', as_attachment=True,attachment_filename="output.pdf")

@app.route('/delete')
def deletepdf ():
    os.remove("F:\\PycharmProjects\\finalproject\\static\\output.pdf")
    return "deleted"

if __name__ == '__main__':
    app.run(port=4545, debug=True)



