import os
import sys
from docx import Document
from docx.shared import Inches
import Word2PDF as pdf
import datetime


def template(imagefile):
    document = Document('F:\\PycharmProjects\\finalproject\\test.docx')
    for file in imagefile:
        target = "<" + (file.filename).split('.')[0] + "|"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target in cell.text:
                        mycell = cell.text
                        arr = (str(mycell)).split('|')
                        cell.text = ""
                        print(arr[2].split('>')[0])
                        print(arr[1])
                        if arr[2].split('>')[0] != "":
                            height = int(arr[2].split('>')[0]) / 96
                        else:
                            height =int(1)
                        if arr[1] != "":
                            width = int(arr[1]) / 96
                        else:
                            width =int(1)
                        run = cell.paragraphs[0].add_run().add_picture(file, width=Inches(height), height=Inches(width))
                        inline_shape = run
                        # document.save('test1.docx')
                        break
                    if "{{date}}" in cell.text:
                        cell.text = ""
                        date=datetime.datetime.now().strftime("%Y-%m-%d")
                        run = cell.add_paragraph(date)

        for para in document.paragraphs:
            if target in para.text:
                mycell = para.text.split('|')
                if mycell[2].split('>')[0]!="":
                    height=int(mycell[2].split('>')[0])/96
                else:
                    height=96
                if mycell[1]!="":
                    width=int(mycell[1])/96
                else:
                    width=96
                para.text = ""
                para.add_run().add_picture(file, width=Inches(height), height=Inches(width))
    document.save('test1.docx')
    pdf.covx_to_pdf("output")

